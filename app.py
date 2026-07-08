import streamlit as st
from docx import Document
import subprocess
import tempfile
import os
import shutil
from datetime import date
from num2words import num2words

TEMPLATE_DOCX = "SampleDoc.docx"
LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"
INVOICE_COUNTER_FILE = "invoice_counter.txt"

def get_today_date():
    return date.today().strftime("%d-%m-%Y")


def amount_to_words(amount):
    try:
        return num2words(amount, lang="en_IN").title() + " Only"
    except Exception:
        return ""


def get_next_invoice_number():
    if not os.path.exists(INVOICE_COUNTER_FILE):
        return "0001"
    try:
        with open(INVOICE_COUNTER_FILE, "r") as f:
            return f"{int(f.read().strip()) + 1:04d}"
    except Exception:
        return "0001"


def save_invoice_number(invoice_number):
    try:
        with open(INVOICE_COUNTER_FILE, "w") as f:
            f.write(f"{int(invoice_number):04d}")
    except ValueError:
        pass


def reset_invoice_counter(next_number):
    with open(INVOICE_COUNTER_FILE, "w") as f:
        f.write(f"{int(next_number) - 1:04d}")


def replace_placeholders(input_docx, output_docx, replacements):
    doc = Document(input_docx)

    def replace_in_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        replaced = False
        for key, value in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, value)
                replaced = True
        if replaced and paragraph.runs:
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = full_text

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    doc.save(output_docx)


def convert_docx_to_pdf(docx_path, output_dir):
    subprocess.run(
        [
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to",
            "pdf",
            docx_path,
            "--outdir",
            output_dir,
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )



st.set_page_config(page_title="Invoice Generator", layout="centered")
st.title(" Invoice Generator")
st.caption("Real-time Preview • Editable Invoice Number • Audit-Safe")

if not os.path.exists(TEMPLATE_DOCX):
    st.error(" SampleDoc.docx not found.")
    st.stop()

if not os.path.exists(LIBREOFFICE_PATH):
    st.error(" LibreOffice not found.")
    st.stop()

with st.expander(" Reset Invoice Number (Advanced)"):
    st.warning("Reset only in exceptional cases.")
    if st.checkbox("I understand the risk"):
        new_start = st.number_input("Set NEXT invoice number to", min_value=1, step=1)
        if st.button(" Reset Counter"):
            reset_invoice_counter(new_start)
            st.success(f"Next invoice will be {new_start:04d}")
            st.rerun()



st.subheader(" Invoice Information")

col1, col2 = st.columns(2)
with col1:
    invoice_no = st.text_input(
        "Invoice Number (Editable)",
        value=get_next_invoice_number()
    )
with col2:
    invoice_date = st.text_input("Invoice Date", get_today_date())

st.subheader("🏢 Buyer Details")
company = st.text_input("Company Name")
addr1 = st.text_input("Address Line 1")
addr2 = st.text_input("Address Line 2")
gst = st.text_input("GST Number")

st.subheader(" Transport Details")
col3, col4 = st.columns(2)
with col3:
    vehicle = st.text_input("Vehicle Number")
    driver = st.text_input("Driver Contact")
with col4:
    supply_place = st.text_input("Place of Supply")
    broker = st.text_input("Broker Name")

st.subheader(" Item Details")
col5, col6, col7 = st.columns(3)
with col5:
    item = st.text_input("Item Name")
    bags = st.text_input("No. of Bags")
with col6:
    weight = st.number_input("Weight (Quintal))", min_value=0.0, step=1.0)
with col7:
    rate = st.number_input("Rate", min_value=0.0, step=1.0)

transport = st.text_input("Transport")


amount = rate * weight
amount_words = amount_to_words(int(round(amount)))

st.info(
    f"""
 **Live Invoice Preview**

**Invoice No:** {invoice_no}  
**Date:** {invoice_date}  
**Item:** {item}  
**Weight:** {weight:.2f} kg  
**Rate:** ₹ {rate:.2f}  
**Total Amount:** ₹ {amount:.2f}  
**Amount in Words:** {amount_words}
"""
)


if st.button("🚀 Generate Invoice"):
    replacements = {
        "<Invoice_Number>": invoice_no,
        "<Date>": invoice_date,
        "<To_Company_Name>": company,
        "<To_Address_1>": addr1,
        "<To_Address_2>": addr2,
        "<GST_NUMBER>": gst,
        "<Vehicle_Number>": vehicle,
        "<DRIVER_CONTACT>": driver,
        "<Supply_Place>": supply_place,
        "<Broker_Name>": broker,
        "<ITEM_NAME>": item,
        "<BAGS>": bags,
        "<WEIGHT>": f"{weight:.2f}",
        "<RATE>": f"{rate:.2f}",
        "<TRANSPORT>": transport,
        "<AMOUNT>": f"{amount:.2f}",
        "<AMOUNT_IN_WORDS>": amount_words,
    }

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "Invoice.docx")
        shutil.copy(TEMPLATE_DOCX, docx_path)
        replace_placeholders(docx_path, docx_path, replacements)
        save_invoice_number(invoice_no)

        st.success(f" Invoice {invoice_no} generated")

        try:
            convert_docx_to_pdf(docx_path, tmpdir)
            pdf_path = docx_path.replace(".docx", ".pdf")
        except Exception:
            pdf_path = None

        with open(docx_path, "rb") as f:
            st.download_button(
                " Download DOCX",
                f,
                file_name=f"Invoice_{invoice_no}.docx"
            )

        if pdf_path and os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                st.download_button(
                    " Download PDF",
                    f,
                    file_name=f"Invoice_{invoice_no}.pdf"
                )

